import os
from docx import Document

# def read_docx(file_path):
#     doc = Document(file_path)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return '\n'.join(full_text)

# file_path = r'E:\python_project\python_docs\python_docs\todo.docx'
# text = read_docx(file_path)
# print(text)


# 888888888888888888888888888888888888

# def search_points_in_docx(file_path, keyword):
#     doc = Document(file_path)
#     points = []
#     for para in doc.paragraphs:
#         if keyword in para.text:
#             points.append(para.text)
#     return points

# file_path = r'E:\python_project\python_docs\python_docs\todo.docx'
# keyword = 'React Project Setup'  # Modify this with the keyword you want to search for
# found_points = search_points_in_docx(file_path, keyword)
# if found_points:
#     print("Points found containing '{}':".format(keyword))
#     for point in found_points:
#         print(point)
# else:
#     print("No points containing '{}' found.".format(keyword))

# =============================================================

# def find_heading(document, heading_text):
#     for paragraph in document.paragraphs:
#         if paragraph.style.name.startswith('Heading'):
#             if paragraph.text == heading_text:
#                 return paragraph

#     return None

# def extract_table_data(file_path):
#     heading_text_to_find = 'HELO'
#     doc = Document(file_path)
#     heading_paragraph = find_heading(doc, heading_text_to_find)
#     if heading_paragraph:
#         print('Heading Found')
#     else:
#         print("Heading No Found")
#     tables_data = []
#     for table in doc.tables:
#         table_data = []
#         for row in table.rows:
#             row_data = []
#             for cell in row.cells:
#                 row_data.append(cell.text.strip())
#             table_data.append(row_data)
#         tables_data.append(table_data)
#     return tables_data

# def print_table_data(tables_data):
#     for i, table_data in enumerate(tables_data):
#         if i+1 == 1:
#             print(f"demo {i+1}:")
#             for row in table_data:
#                 print("\t".join(row))
#             print()
            
#         else:
#             print(f"Table {i+1}:")
#             for row in table_data:
#                 print("\t".join(row))
#             print()
        


# file_path = r'E:\python_project\python_docs\python_docs\demo.docx'
# filename = os.path.basename(file_path)
# print("Filename:", filename)
# tables_data = extract_table_data(file_path)
# print_table_data(tables_data)
# ++++++++++++++++++++++++++++++++++++++++++++++++

# def extract_table_column_names(file_path, keyword):
#     doc = Document(file_path)
#     for para in doc.paragraphs:
#         if keyword in para.text:
#             column_names = []
#             for table in doc.tables:
#                 if len(table.rows) > 0:
#                     first_row = table.rows[0]
#                     for cell in first_row.cells:
#                         column_names.append(cell.text.strip())
#                     break
#             return column_names

# file_path = r'E:\python_project\python_docs\python_docs\todo_table.docx'
# filename = os.path.basename(file_path)
# print("Filename>>>>>>>>>>>>>>>:", filename)
# first_value = filename.split('_')[0]
# print('first_value>>>>', first_value)
# keyword = 'Applying K-means with no. clusters = 7'
# column_names = extract_table_column_names(file_path, keyword)
# print("Table Column Names:", column_names)

# from docx import Document



# import pandas as pd

# ============Table Heading FOund code =============================



# ================================================================




# import pandas as pd
# from docx import Document
# document = Document(r'E:\python_project\python_docs\python_docs\todo_table.docx')
# print('-----------------')

# data = [[cell.text for cell in row.cells] for row in table.rows]
# df = pd.DataFrame(data)

# df = df.rename(columns=df.iloc[0]).drop(df.index[0]).reset_index(drop=True)


# outside_col, inside_col = df.iloc[0], df.iloc[1]
# hier_index = pd.MultiIndex.from_tuples(list(zip(outside_col, inside_col)))
# df = pd.DataFrame(data,columns=hier_index).drop(df.index[[0,1]] ).reset_index(drop=True)



# from docx import Document
# file_path1 = r'E:\python_project\python_docs\python_docs\todo_table.docx'
# wordDoc = Document(file_path1)

# for table in wordDoc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print(cell.text)

# def read_docx_table(file_path, table_heading):
#     # Load the .docx file
#     doc = Document(file_path)
    
#     # Initialize variables
#     table_found = False
#     data = []

#     # Iterate through paragraphs
#     for para in doc.paragraphs:
#         print('>>>>>>>>>>>>>>>>', table_heading)
#         # Check if paragraph text matches the specified heading
#         if para.text.strip() == table_heading:
#             table_found = True
#         # If the table is found, read the table
#         elif table_found:
#             # Check if the paragraph contains a table
#             if para.tables:
#                 # Iterate through tables
#                 for table in para.tables:
#                     # Convert the table to a pandas DataFrame
#                     df = pd.DataFrame([[cell.text for cell in row.cells] for row in table.rows])
#                     # Append the DataFrame to the data list
#                     data.append(df)
#             # Break the loop if we have processed the table
#             break

#     # Concatenate DataFrames if there are multiple tables
#     if len(data) > 1:
#         return pd.concat(data)
#     elif len(data) == 1:
#         return data[0]
#     else:
#         return None

# # Example usage
# file_path = r'E:\python_project\python_docs\python_docs\SampleDOCFile_100kb.docx'
# table_heading = "Here is a table"
# table_data = read_docx_table(file_path, table_heading)
# print(table_data)

