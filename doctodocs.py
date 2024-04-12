# from docx import Document

# def convert_doc_to_docx(doc_file, docx_file):
#     # Open the .doc file
#     with open(doc_file, 'rb') as doc:
#         docx = Document(doc)
    
#     # Save as .docx file
#     docx.save(docx_file)

# # Provide the paths of the input .doc file and the output .docx file
# input_doc_file = 'input.doc'
# output_docx_file = 'output.docx'

# convert_doc_to_docx(input_doc_file, output_docx_file)


from docx import Document

def extract_table_headings(docx_file_path, table_index=0):
    doc = Document(docx_file_path)
    table = doc.tables[table_index]
    
    headings = []
    for cell in table.rows[0].cells:
        headings.append(cell.text.strip())
    
    return headings

# Example usage:
docx_file_path = "your_document.docx"
headings = extract_table_headings(docx_file_path)
print("Table Headings:", headings)